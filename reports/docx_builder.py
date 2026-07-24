"""Generación de reportes Word (.docx) a partir de los resultados calculados.

Usa python-docx puro (sin LibreOffice), para que funcione igual en local y en
Streamlit Community Cloud. El reporte de voladura replica el esquema de un
informe técnico de perforación y voladura real: encabezado del proyecto,
introducción, programa de actividades por etapa, una sección numerada por
labor (diseño de trazo, aspectos técnicos, cálculo de avance/volumen/
tonelaje, explosivos, accesorios, cuadro resumen, memoria de cálculo), y los
cuadros consolidados del programa completo al final.
"""

from __future__ import annotations

import datetime as _dt
from collections import defaultdict
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from core.constants import ACTIVIDADES_CICLICAS, DESCRIPCION_TIPO_LABOR, PROPOSITO_TIPO_LABOR
from core.memoria import memoria_calculo
from core.models import (
    DatosGenerales,
    LaborMinera,
    Polvorin,
    PuntoRiesgo,
    ResultadoDistancia,
    ResultadoVoladura,
)
from core.polvorin import area_shoelace, perimetro
from reports.narrativa import parrafos_introduccion, programa_actividades, secuencia_operativa
from viz.tunnel_plot import build_tunnel_figure


def _fmt(value, decimals=2) -> str:
    if isinstance(value, bool):
        return str(value)
    if isinstance(value, float):
        return f"{value:,.{decimals}f}"
    if isinstance(value, int):
        return f"{value:,}"
    return str(value)


def _add_table(document: Document, headers: list[str], rows: list[list]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = _fmt(val) if not isinstance(val, str) else val


def _add_title_page(document: Document, titulo: str, subtitulo: str = "") -> None:
    h = document.add_heading(titulo, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitulo:
        p = document.add_paragraph(subtitulo)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha = document.add_paragraph(f"Generado el {_dt.date.today().isoformat()}")
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("")


def _add_encabezado_informe(document: Document, datos: DatosGenerales, titulo_proyecto: str) -> None:
    """Portada del informe técnico: nombre de la concesión/proyecto, código,
    empresa, RUC, ubicación y fecha — replica el bloque de encabezado de un
    informe técnico de perforación y voladura. Cada línea se omite si el
    dato correspondiente no fue consignado (nunca se inventa un valor)."""
    titulo_principal = datos.nombre_concesion or titulo_proyecto
    h = document.add_heading(titulo_principal.upper(), level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for texto in (datos.codigo_concesion and f"Código: {datos.codigo_concesion}", datos.empresa, datos.ruc and f"RUC: {datos.ruc}"):
        if texto:
            p = document.add_paragraph(texto)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    h2 = document.add_heading("INFORME TECNICO", level=1)
    h2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if datos.nombre_concesion and titulo_proyecto:
        p = document.add_paragraph(titulo_proyecto)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ubicacion = ", ".join(filter(None, [datos.distrito, datos.provincia, datos.departamento]))
    if ubicacion:
        p = document.add_paragraph(ubicacion.upper())
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fecha = document.add_paragraph(f"Generado el {_dt.date.today().isoformat()}")
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("")


def _add_labor_section(
    document: Document, numero: int, labor: LaborMinera, resultado: ResultadoVoladura
) -> None:
    document.add_heading(f"{numero}. {labor.tipo.upper()}: {labor.nombre}", level=1)

    descripcion = DESCRIPCION_TIPO_LABOR.get(labor.tipo)
    if descripcion:
        document.add_paragraph(descripcion)
    proposito = PROPOSITO_TIPO_LABOR.get(labor.tipo, "de acuerdo con el diseño de perforación adoptado")
    document.add_paragraph(
        f"Se ha programado el desarrollo de {labor.nombre}, con sección de "
        f"{labor.ancho_m:.2f} m × {labor.alto_m:.2f} m, {proposito}."
    )
    if labor.observaciones:
        document.add_paragraph(labor.observaciones)

    try:
        png_bytes = build_tunnel_figure(labor, resultado).to_image(
            format="png", width=900, height=550, scale=2
        )
        document.add_picture(BytesIO(png_bytes), width=Inches(6))
    except Exception:
        document.add_paragraph("(Esquema 3D no disponible en este entorno)")

    document.add_heading(f"{numero}.1 Diseño del Trazo", level=2)
    document.add_paragraph(
        f"La labor {labor.nombre} se ejecutará mediante perforación y "
        f"voladura convencional, empleando equipo tipo "
        f"{labor.equipo_perforacion}, con los siguientes parámetros técnicos:"
    )
    _add_table(
        document,
        ["Concepto", "Valor"],
        [
            ["Labor", labor.nombre],
            ["Sección", f"{labor.ancho_m} × {labor.alto_m} m"],
            ["Área de la sección", f"{_fmt(resultado.area_m2, 3)} m²"],
            ["Longitud existente", f"{_fmt(labor.longitud_existente_m)} m"],
            ["Avance proyectado", f"{_fmt(labor.avance_proyectado_m)} m"],
            ["Longitud final", f"{_fmt(resultado.longitud_final_m)} m"],
            ["Longitud del barreno", f"{labor.longitud_barreno_pies} pies"],
            ["Diámetro del barreno", f"{labor.diametro_barreno_mm} mm"],
            ["Equipo de perforación", labor.equipo_perforacion],
            ["Tipo de corte", labor.tipo_corte],
            ["Taladros cargados por disparo", labor.taladros_cargados],
            ["Taladros de alivio", labor.taladros_alivio],
        ],
    )

    document.add_heading(f"{numero}.2 Aspectos Técnicos", level=2)
    _add_table(
        document,
        ["Concepto", "Valor"],
        [
            ["Tipo de roca", labor.tipo_roca],
            ["Destino del material", labor.destino_material],
            ["Peso específico usado", f"{_fmt(resultado.densidad_usada_tm_m3)} TM/m³"],
            ["Método de avance", "Perforación y voladura convencional"],
        ],
    )

    document.add_heading(f"{numero}.3 Cálculo de Avance, Volumen y Tonelaje", level=2)
    _add_table(
        document,
        ["Concepto", "Valor"],
        [
            ["Número de disparos", resultado.n_disparos],
            ["Volumen por disparo", f"{_fmt(resultado.volumen_por_disparo_m3)} m³"],
            ["Volumen total", f"{_fmt(resultado.volumen_total_m3)} m³"],
            ["Tonelaje por disparo", f"{_fmt(resultado.tonelaje_por_disparo_tm)} TM"],
            ["Tonelaje total", f"{_fmt(resultado.tonelaje_total_tm)} TM"],
        ],
    )

    document.add_heading(f"{numero}.4 Explosivos para la Voladura", level=2)
    document.add_paragraph(
        f"Para el avance de {labor.nombre} se empleará {labor.tipo_explosivo_1} "
        f"en los taladros de arranque y ayuda, mientras que {labor.tipo_explosivo_2} "
        f"será utilizada como carga principal en los taladros de producción y "
        f"contorno, de acuerdo con la distribución de la malla y las "
        f"características de la roca {labor.tipo_roca.lower()}."
    )
    _add_table(
        document,
        ["Concepto", "Valor"],
        [
            ["Cartuchos por taladro", labor.cartuchos_por_taladro],
            ["Cartuchos por disparo", resultado.cartuchos_por_disparo],
            ["Explosivo por disparo", f"{_fmt(resultado.explosivo_por_disparo_kg)} kg"],
            ["Explosivo total", f"{_fmt(resultado.explosivo_total_kg)} kg"],
            [f"{labor.tipo_explosivo_1} ({labor.pct_explosivo_1:.0f}%)", f"{_fmt(resultado.explosivo_tipo1_kg)} kg"],
            [f"{labor.tipo_explosivo_2} ({labor.pct_explosivo_2:.0f}%)", f"{_fmt(resultado.explosivo_tipo2_kg)} kg"],
            ["Factor de potencia", f"{_fmt(resultado.factor_potencia_kg_tm)} kg/TM"],
            ["Consumo específico por volumen", f"{_fmt(resultado.consumo_especifico_kg_m3)} kg/m³"],
        ],
    )

    document.add_heading(f"{numero}.5 Accesorios para la Voladura", level=2)
    document.add_paragraph(
        f"Se considera un {labor.tipo_fulminante.lower()} y un tramo de mecha "
        f"de seguridad por cada taladro cargado. Los {labor.taladros_alivio} "
        f"taladros de alivio permanecerán sin carga."
    )
    _add_table(
        document,
        ["Concepto", "Valor"],
        [
            [labor.tipo_fulminante, f"{resultado.fulminantes_total} unidades"],
            ["Mecha de seguridad por taladro", f"{_fmt(resultado.mecha_por_taladro_m, 3)} m"],
            ["Mecha de seguridad por disparo", f"{_fmt(resultado.mecha_por_disparo_m)} m"],
            ["Mecha de seguridad total", f"{_fmt(resultado.mecha_total_m)} m"],
        ],
    )

    document.add_heading(f"Cuadro Resumen — {labor.nombre}", level=2)
    _add_table(
        document,
        ["Concepto", "Valor"],
        [
            ["Sección", f"{labor.ancho_m} × {labor.alto_m} m"],
            ["Área", f"{_fmt(resultado.area_m2, 3)} m²"],
            ["Avance proyectado", f"{_fmt(labor.avance_proyectado_m)} m"],
            ["Disparos", resultado.n_disparos],
            ["Taladros cargados/disparo", labor.taladros_cargados],
            ["Cartuchos/disparo", resultado.cartuchos_por_disparo],
            ["Explosivo/disparo", f"{_fmt(resultado.explosivo_por_disparo_kg)} kg"],
            ["Explosivo total", f"{_fmt(resultado.explosivo_total_kg)} kg"],
            [labor.tipo_explosivo_1, f"{_fmt(resultado.explosivo_tipo1_kg)} kg"],
            [labor.tipo_explosivo_2, f"{_fmt(resultado.explosivo_tipo2_kg)} kg"],
            [labor.tipo_fulminante, f"{resultado.fulminantes_total} unidades"],
            ["Mecha de seguridad", f"{_fmt(resultado.mecha_total_m)} m"],
            ["Volumen total", f"{_fmt(resultado.volumen_total_m3)} m³"],
            ["Tonelaje total", f"{_fmt(resultado.tonelaje_total_tm)} TM"],
            ["TM/disparo", f"{_fmt(resultado.tonelaje_por_disparo_tm)} TM"],
            ["Factor de potencia", f"{_fmt(resultado.factor_potencia_kg_tm)} kg/TM"],
        ],
    )

    document.add_heading("Memoria de cálculo", level=2)
    document.add_paragraph(
        "Desglose paso a paso (fórmula → sustitución → resultado) de cada "
        "cifra reportada arriba."
    )
    pasos = memoria_calculo(labor, resultado)
    _add_table(
        document,
        ["Concepto", "Fórmula", "Sustitución", "Resultado"],
        [[p.concepto, p.formula, p.sustitucion, p.resultado] for p in pasos],
    )
    document.add_paragraph("")


def _add_cuadros_consolidados(
    document: Document, labores: list[LaborMinera], resultados: list[ResultadoVoladura]
) -> None:
    document.add_heading("Cuadro Resumen General del Programa", level=1)
    filas = [
        [
            labor.nombre,
            f"{labor.ancho_m} × {labor.alto_m} = {resultado.area_m2:.2f} m²",
            labor.etapa,
            labor.destino_material,
            f"{_fmt(resultado.tonelaje_total_tm)} TM",
            f"{labor.taladros_cargados} / {_fmt(labor.avance_proyectado_m)} m",
        ]
        for labor, resultado in zip(labores, resultados)
    ]
    total_tonelaje = sum(r.tonelaje_total_tm for r in resultados)
    filas.append(["TOTAL", "—", "—", "—", f"{_fmt(total_tonelaje)} TM", "—"])
    _add_table(
        document,
        ["Labor minera", "Especificación técnica", "Etapa", "Material", "Tonelaje", "Taladros / Avance"],
        filas,
    )

    document.add_heading("Tonelaje por Disparo", level=1)
    _add_table(
        document,
        ["Labor", "TM/disparo"],
        [[labor.nombre, f"{_fmt(r.tonelaje_por_disparo_tm)}"] for labor, r in zip(labores, resultados)],
    )

    document.add_heading("Explosivo por Disparo", level=1)
    _add_table(
        document,
        ["Labor", "Kg/disparo"],
        [[labor.nombre, f"{_fmt(r.explosivo_por_disparo_kg)} kg"] for labor, r in zip(labores, resultados)],
    )

    document.add_heading("Número de Disparos por Labor", level=1)
    filas_disparos = [[labor.nombre, r.n_disparos] for labor, r in zip(labores, resultados)]
    filas_disparos.append(["TOTAL", sum(r.n_disparos for r in resultados)])
    _add_table(document, ["Labor", "Disparos"], filas_disparos)

    document.add_heading("Resumen de Avance y Explosivos", level=1)
    filas_resumen = [
        [
            labor.nombre,
            f"{_fmt(labor.avance_proyectado_m)} m",
            r.n_disparos,
            f"{_fmt(r.explosivo_por_disparo_kg)} kg",
            f"{_fmt(r.explosivo_total_kg)} kg",
        ]
        for labor, r in zip(labores, resultados)
    ]
    filas_resumen.append(
        [
            "TOTAL GENERAL",
            f"{_fmt(sum(l.avance_proyectado_m for l in labores))} m",
            sum(r.n_disparos for r in resultados),
            "—",
            f"{_fmt(sum(r.explosivo_total_kg for r in resultados))} kg",
        ]
    )
    _add_table(
        document,
        ["Labor", "Avance proyectado", "N.° disparos", "Kg por disparo", "Total de explosivos"],
        filas_resumen,
    )

    document.add_heading("Distribución de Explosivos", level=1)
    por_explosivo: dict[str, float] = defaultdict(float)
    for labor, r in zip(labores, resultados):
        por_explosivo[labor.tipo_explosivo_1] += r.explosivo_tipo1_kg
        por_explosivo[labor.tipo_explosivo_2] += r.explosivo_tipo2_kg
    total_explosivo = sum(por_explosivo.values())
    filas_exp = [
        [nombre, f"{_fmt(kg)} kg", f"{(kg / total_explosivo * 100) if total_explosivo else 0:.1f} %"]
        for nombre, kg in por_explosivo.items()
    ]
    filas_exp.append(["TOTAL GENERAL", f"{_fmt(total_explosivo)} kg", "100.0 %"])
    _add_table(document, ["Explosivo", "Total", "Porcentaje"], filas_exp)

    document.add_heading("Fulminantes y Mecha por Labor", level=1)
    filas_acc = [
        [labor.nombre, r.n_disparos, f"{r.fulminantes_total} und", f"{_fmt(r.mecha_total_m)} m"]
        for labor, r in zip(labores, resultados)
    ]
    filas_acc.append(
        [
            "TOTAL GENERAL",
            sum(r.n_disparos for r in resultados),
            f"{sum(r.fulminantes_total for r in resultados)} und",
            f"{_fmt(sum(r.mecha_total_m for r in resultados))} m",
        ]
    )
    _add_table(document, ["Labor", "N.° de disparos", "Fulminantes", "Mecha de seguridad"], filas_acc)

    document.add_heading("Totales de Accesorios", level=1)
    _add_table(
        document,
        ["Accesorio", "Total"],
        [
            ["Mecha de seguridad", f"{_fmt(sum(r.mecha_total_m for r in resultados))} m"],
            ["Fulminantes", f"{sum(r.fulminantes_total for r in resultados)} unidades"],
        ],
    )


def build_voladura_report(
    labores: list[LaborMinera],
    resultados: list[ResultadoVoladura],
    titulo_proyecto: str = "Programa de perforación y voladura",
    datos: DatosGenerales | None = None,
) -> BytesIO:
    datos = datos or DatosGenerales()
    document = Document()

    _add_encabezado_informe(document, datos, titulo_proyecto)
    document.add_heading("PROGRAMA DE AVANCES DE LABORES", level=1)

    document.add_heading("Introducción", level=1)
    for parrafo in parrafos_introduccion(labores, resultados, datos):
        document.add_paragraph(parrafo)

    if datos.distrito or datos.provincia or datos.departamento:
        document.add_heading("Ubicación Política", level=1)
        nombre = datos.nombre_concesion or "El proyecto"
        ubicacion = ", ".join(
            filter(None, [
                datos.distrito and f"distrito de {datos.distrito}",
                datos.provincia and f"provincia de {datos.provincia}",
                datos.departamento and f"departamento de {datos.departamento}",
            ])
        )
        document.add_paragraph(f"{nombre} se ubica en el {ubicacion}.")

    document.add_heading("Actividades Cíclicas de Actividades Mineras", level=1)
    document.add_paragraph("Las actividades cíclicas serán las siguientes:")
    for actividad in ACTIVIDADES_CICLICAS:
        document.add_paragraph(actividad, style="List Bullet")

    document.add_heading("Programa de Actividades", level=1)
    document.add_paragraph(
        "Las labores mineras se ejecutarán de manera secuencial, "
        "comprendiendo las etapas de desarrollo, preparación y explotación, "
        "conforme al programa establecido."
    )
    for etapa, intro, bullets in programa_actividades(labores, resultados):
        document.add_heading(f"Etapa de {etapa}", level=2)
        if intro:
            document.add_paragraph(intro)
        for bullet in bullets:
            document.add_paragraph(bullet, style="List Bullet")

    document.add_heading("Secuencia Operativa", level=1)
    for linea in secuencia_operativa(labores, resultados, datos):
        document.add_paragraph(linea)

    document.add_heading("Especificaciones Técnicas del Programa y Diseño de Voladura", level=1)
    for i, (labor, resultado) in enumerate(zip(labores, resultados), start=1):
        _add_labor_section(document, i, labor, resultado)

    _add_cuadros_consolidados(document, labores, resultados)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def build_polvorin_report(
    polvorines: list[Polvorin],
    resultados_por_polvorin: dict[str, list[ResultadoDistancia]],
    titulo_proyecto: str = "Verificación de seguridad de polvorín",
) -> BytesIO:
    document = Document()
    _add_title_page(document, titulo_proyecto, "Reporte generado automáticamente")
    document.add_paragraph(
        "Nota: las distancias mínimas requeridas mostradas en este reporte son "
        "las capturadas manualmente en la aplicación. Deben verificarse contra "
        "el reglamento vigente (D.S. N.° 024-2016-EM y modificatorias) antes de "
        "tomar decisiones operativas."
    )

    for polvorin in polvorines:
        document.add_heading(f"Polvorín de {polvorin.tipo}: {polvorin.nombre}", level=1)
        _add_table(
            document,
            ["Concepto", "Valor"],
            [
                ["Coordenadas (Este, Norte UTM)", f"{polvorin.este_utm}, {polvorin.norte_utm}"],
                [
                    "Área del cerco perimétrico",
                    f"{_fmt(area_shoelace(polvorin.vertices_cerco))} m²"
                    if polvorin.vertices_cerco
                    else "No consignada",
                ],
                [
                    "Perímetro del cerco",
                    f"{_fmt(perimetro(polvorin.vertices_cerco))} m"
                    if polvorin.vertices_cerco
                    else "No consignado",
                ],
                [
                    "Cantidad almacenada",
                    f"{_fmt(polvorin.cantidad_almacenada_kg)} kg"
                    if polvorin.cantidad_almacenada_kg
                    else "No consignada",
                ],
                [
                    "Radio de influencia",
                    f"{_fmt(polvorin.radio_influencia_m)} m"
                    if polvorin.radio_influencia_m
                    else "No consignado",
                ],
            ],
        )

        resultados = resultados_por_polvorin.get(polvorin.nombre, [])
        if resultados:
            document.add_heading("Distancias a puntos de riesgo", level=2)
            _add_table(
                document,
                [
                    "Punto de riesgo",
                    "Tipo",
                    "Distancia real (m)",
                    "Distancia mínima requerida (m)",
                    "¿Cumple?",
                ],
                [
                    [
                        r.punto_nombre,
                        r.punto_tipo,
                        _fmt(r.distancia_real_m),
                        _fmt(r.distancia_minima_m),
                        "Sí" if r.cumple else "No",
                    ]
                    for r in resultados
                ],
            )
        document.add_paragraph("")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
